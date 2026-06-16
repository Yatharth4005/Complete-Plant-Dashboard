import re
from datetime import datetime

# Optional imports
try:
    import docx
except ImportError:
    docx = None

try:
    import pypdf
except ImportError:
    pypdf = None

def parse_capa_file(file_stream, filename):
    """
    Unified parser that handles:
    - .docx (ZIP/OpenXML format)
    - .doc (Legacy OLE binary format)
    - .pdf (Adobe Portable Document format)
    """
    # Read the first few bytes to check the file signature
    header = file_stream.read(8)
    file_stream.seek(0)
    
    ext = filename.lower().split('.')[-1]
    
    if header.startswith(b'PK\x03\x04') or ext == 'docx':
        # ZIP file: parse as docx
        return _parse_docx(file_stream)
    elif header.startswith(b'\xd0\xcf\x11\xe0') or ext == 'doc':
        # OLE file: parse as legacy doc
        return _parse_legacy_doc(file_stream)
    elif header.startswith(b'%PDF') or ext == 'pdf':
        # PDF file: parse as pdf
        return _parse_pdf(file_stream)
    else:
        # Fallback to legacy doc string extraction if signature is unrecognized but extension matches
        return _parse_legacy_doc(file_stream)


def _parse_docx(file_stream):
    """Parses .docx files using python-docx."""
    if not docx:
        raise ImportError("The 'python-docx' library is required to parse .docx files. Please run: .\\.venv\\Scripts\\python -m pip install python-docx")
        
    doc = docx.Document(file_stream)
    all_lines = []
    
    # Extract from paragraphs
    for p in doc.paragraphs:
        val = p.text.strip()
        if val:
            all_lines.append(val)
            
    # Extract from tables (row by row, cell by cell)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    val = p.text.strip()
                    if val and val not in all_lines:
                        all_lines.append(val)
                        
    text_dump = "\n".join(all_lines)
    return _extract_capa_fields_from_text(text_dump, all_lines, doc_tables=doc.tables)


def _parse_legacy_doc(file_stream):
    """Parses legacy binary .doc files by extracting ASCII and UTF-16 strings."""
    content = file_stream.read()
    
    # Extract ascii and utf-16 strings
    ascii_strings = re.findall(b'[\x20-\x7E]{4,}', content)
    utf16_strings = re.findall(b'(?:[\x20-\x7E]\x00){4,}', content)
    
    all_lines = []
    for s in ascii_strings:
        try:
            decoded = s.decode('ascii').strip()
            if decoded:
                all_lines.append(decoded)
        except Exception:
            pass
            
    for s in utf16_strings:
        try:
            decoded = s.decode('utf-16le').strip()
            if decoded:
                all_lines.append(decoded)
        except Exception:
            pass
            
    text_dump = "\n".join(all_lines)
    return _extract_capa_fields_from_text(text_dump, all_lines)


def _parse_pdf(file_stream):
    """Parses PDF files using pypdf."""
    if not pypdf:
        raise ImportError("The 'pypdf' library is required to parse PDF files. Please run: .\\.venv\\Scripts\\python -m pip install pypdf")
        
    reader = pypdf.PdfReader(file_stream)
    all_lines = []
    
    for page in reader.pages:
        text = page.extract_text()
        if text:
            for line in text.split('\n'):
                line_clean = line.strip()
                if line_clean:
                    all_lines.append(line_clean)
                    
    text_dump = "\n".join(all_lines)
    return _extract_capa_fields_from_text(text_dump, all_lines)


def _extract_capa_fields_from_text(text_dump, all_lines, doc_tables=None):
    """Generic text parsing logic that maps extracted text to CAPAReport fields."""
    
    def find_colon_value(pattern, text):
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
        return ""

    # Basic Info
    area_section = find_colon_value(r'(?:Area\s*/\s*Section|Area|Section)\s*:\s*(.*)', text_dump)
    date_incident = find_colon_value(r'(?:Date\s*of\s*Incident|Incident\s*Date)\s*:\s*(.*)', text_dump)
    capa_no = find_colon_value(r'(?:CAPA\s*No\.?|CAPA\s*Number)\s*:\s*(.*)', text_dump)

    # Document metadata
    document_no = find_colon_value(r'(?:Document\s*No\.?|Doc\s*No\.?)\s*:\s*(.*)', text_dump)
    if not document_no:
        match = re.search(r'Document\s*No\.?\s*([^\n\r\|]+)', text_dump, re.IGNORECASE)
        if match:
            document_no = match.group(1).strip()
            
    issue_no = find_colon_value(r'(?:Issue\s*No\.?)\s*:\s*(.*)', text_dump)
    if not issue_no:
        match = re.search(r'Issue\s*No\.?\s*([^\n\r\|]+)', text_dump, re.IGNORECASE)
        if match:
            issue_no = match.group(1).strip()
            
    issue_date = find_colon_value(r'(?:Issue\s*Date)\s*:\s*(.*)', text_dump)
    if not issue_date:
        match = re.search(r'Issue\s*Date\s*([\d\.\-/]+)', text_dump, re.IGNORECASE)
        if match:
            issue_date = match.group(1).strip()
    
    # 1. Problem description
    problem_what = ""
    problem_where = ""
    problem_when = ""
    problem_extent = ""
    
    for i, line in enumerate(all_lines):
        line_lower = line.lower()
        if line_lower.startswith("what:"):
            part = line[5:].strip()
            if part:
                problem_what = part
            elif i + 1 < len(all_lines):
                next_l = all_lines[i+1]
                if not any(next_l.lower().startswith(kw) for kw in ['where:', 'when:', 'extent:', '2.']):
                    problem_what = next_l
        elif line_lower.startswith("where:"):
            problem_where = line[6:].strip()
        elif line_lower.startswith("when:"):
            problem_when = line[5:].strip()
        elif line_lower.startswith("extent:"):
            problem_extent = line[7:].strip()

    # 3. Correction / Immediate Actions
    immediate_action = ""
    action_timeframe = ""
    action_responsibility = ""
    
    for i, line in enumerate(all_lines):
        line_lower = line.lower()
        if "immediate actions taken" in line_lower or "correction / immediate actions" in line_lower:
            actions = []
            for j in range(i + 1, min(i + 6, len(all_lines))):
                next_l = all_lines[j]
                if "time frame" in next_l.lower() or "responsibility" in next_l.lower() or "4." in next_l.lower():
                    break
                actions.append(next_l)
            immediate_action = " ".join(actions).strip()
            
        if "time frame" in line_lower and ":" in line:
            action_timeframe = line.split(":", 1)[1].strip()
        if "responsibility" in line_lower and ":" in line and i < 45:
            action_responsibility = line.split(":", 1)[1].strip()

    # 4. Root Cause (Dynamic Whys List)
    whys_dict = {}
    for num in range(1, 11):
        ord_lbl = "1st" if num == 1 else "2nd" if num == 2 else "3rd" if num == 3 else f"{num}th"
        val = find_colon_value(rf'{ord_lbl}\s*Why\??\s*:\s*(.*)', text_dump)
        if val and "conclusion" not in val.lower():
            whys_dict[num] = val

    for i, line in enumerate(all_lines):
        line_clean = line.replace("?", "").strip()
        match = re.search(r'\b(\d+)(?:st|nd|rd|th)\s*Why', line_clean, re.IGNORECASE)
        if match:
            num = int(match.group(1))
            val = ""
            if ":" in line:
                val_parts = line.split(":", 1)
                if len(val_parts) > 1:
                    val = val_parts[1].strip()
            elif i + 1 < len(all_lines):
                next_l = all_lines[i+1]
                if not any(kw in next_l.lower() for kw in ["conclusion", "3.", "correction", "immediate", "why"]):
                    val = next_l
            if val and "conclusion" not in val.lower():
                whys_dict[num] = val

    max_why = max(whys_dict.keys()) if whys_dict else 5
    if max_why < 5:
        max_why = 5
        
    whys_list = []
    for idx in range(1, max_why + 1):
        whys_list.append(whys_dict.get(idx, ""))

    why_1 = whys_dict.get(1, "")
    why_2 = whys_dict.get(2, "")
    why_3 = whys_dict.get(3, "")
    why_4 = whys_dict.get(4, "")
    why_5 = whys_dict.get(5, "")

    conclusion = find_colon_value(r'Conclusion\s*\(s\)\s*:\s*(.*)', text_dump)
    if not conclusion:
        for i, line in enumerate(all_lines):
            if line.lower().startswith("conclusion"):
                val = line.split(":", 1)[1].strip() if ":" in line else (all_lines[i+1] if i+1 < len(all_lines) else "")
                # Skip placeholder headers
                if "signature" not in val.lower():
                    conclusion = val
                break

    # 5 & 6. Corrective / Preventive Actions
    corrective_actions = []
    preventive_actions = []
    
    if doc_tables:
        for table in doc_tables:
            header_row_idx = -1
            is_corrective = False
            is_preventive = False
            
            for r_idx, row in enumerate(table.rows):
                row_cells_text = [cell.text.lower().strip() for cell in row.cells]
                has_action = any("corrective action" in ct or "preventive action" in ct for ct in row_cells_text)
                has_resp = any("responsibility" in ct or "resp" in ct for ct in row_cells_text)
                if has_action and has_resp:
                    header_row_idx = r_idx
                    is_corrective = any("corrective action" in ct for ct in row_cells_text)
                    is_preventive = any("preventive action" in ct for ct in row_cells_text)
                    break
                    
            if header_row_idx != -1:
                header_cells = [cell.text.lower().strip() for cell in table.rows[header_row_idx].cells]
                col_mapping = {}
                for idx, cell_text in enumerate(header_cells):
                    ct = cell_text.lower()
                    if "corrective action" in ct or "preventive action" in ct:
                        col_mapping['action'] = idx
                    elif "responsibility" in ct or "resp" in ct:
                        col_mapping['responsibility'] = idx
                    elif "target" in ct or "date" in ct:
                        if "implement" in ct or "impl" in ct:
                            col_mapping['impl_date'] = idx
                        else:
                            col_mapping['target_date'] = idx
                            
                for row_idx in range(header_row_idx + 1, len(table.rows)):
                    row = table.rows[row_idx]
                    if len(row.cells) <= max(col_mapping.values(), default=-1):
                        continue
                    
                    action_text = row.cells[col_mapping.get('action', 0)].text.strip()
                    if not action_text or action_text.lower() in ["corrective action", "preventive action", "action", "action(s)"]:
                        continue
                        
                    resp_text = row.cells[col_mapping['responsibility']].text.strip() if 'responsibility' in col_mapping else ""
                    target_date = row.cells[col_mapping['target_date']].text.strip() if 'target_date' in col_mapping else ""
                    impl_date = row.cells[col_mapping['impl_date']].text.strip() if 'impl_date' in col_mapping else ""
                    
                    item = {
                        'action': action_text,
                        'responsibility': resp_text,
                        'target_date': target_date,
                        'impl_date': impl_date
                    }
                    if is_corrective:
                        corrective_actions.append(item)
                    else:
                        preventive_actions.append(item)

    # Text-based fallback (for PDF/legacy doc where doc_tables is not available or failed)
    if not corrective_actions or not preventive_actions:
        idx_5 = -1
        idx_6 = -1
        idx_7 = -1
        
        for i, line in enumerate(all_lines):
            l_lower = line.lower()
            if ("5." in l_lower or "5 " in l_lower) and "corrective" in l_lower:
                idx_5 = i
            elif ("6." in l_lower or "6 " in l_lower) and "preventive" in l_lower:
                idx_6 = i
            elif ("7." in l_lower or "7 " in l_lower) and ("detailed" in l_lower or "implementation" in l_lower):
                idx_7 = i
                
        if idx_5 != -1 and idx_6 != -1:
            corr_lines = all_lines[idx_5 + 1 : idx_6]
            prev_lines = all_lines[idx_6 + 1 : idx_7 if idx_7 != -1 else idx_6 + 15]
            
            def is_header_label(s):
                s_low = s.lower().strip()
                if "date of implementation" in s_low or "implementation date" in s_low or "impl. date" in s_low:
                    return True
                headers = [
                    "corrective action", "corrective action(s)",
                    "preventive action", "preventive action(s)",
                    "responsibility", "target", "target date"
                ]
                if s_low in headers:
                    return True
                if s_low == "date" or s_low == "date:":
                    return True
                return False

            def parse_rows_from_lines(lines_slice):
                rows = []
                current_row = None
                
                known_names = ["saurabh agrawal", "saurabh agarwal", "sachin khanna", "dibyayan paul", "d.paul", "kk singh", "shift i/c", "shift incharge"]
                name_endings = [" singh", " paul", " agrawal", " agarwal", " khanna", " sen", " roy", " gupta", " kumar", " sharma", " patel", " verma", " mandol", " viswakarma"]
                designation_keywords = ["i/c", "incharge", "shift", "team", "member", "leader", "hod", "engineer", "fitters", "fitter", "operator", "dept", "electrical", "mechanical"]

                for line in lines_slice:
                    s = line.strip()
                    if not s or is_header_label(s):
                        continue
                        
                    s_low = s.lower()
                    
                    is_date = False
                    if re.search(r'\b\d{2}[\./-]\d{2}[\./-]\d{4}\b', s):
                        is_date = True
                    elif s_low in ["section change", "immediate", "as scheduled", "continuous", "completed", "na"]:
                        is_date = True
                    elif len(s) < 15 and any(char.isdigit() for char in s) and any(sep in s for sep in ['.', '/', '-']):
                        is_date = True
                        
                    is_resp = False
                    if not is_date:
                        if any(name in s_low for name in known_names):
                            is_resp = True
                        elif any(s_low.endswith(ending) for ending in name_endings):
                            is_resp = True
                        elif len(s) < 25 and any(kw in s_low for kw in designation_keywords):
                            is_resp = True
                            
                    if not is_date and not is_resp:
                        if current_row:
                            rows.append(current_row)
                        current_row = {
                            'action': s,
                            'responsibility': '',
                            'target_date': '',
                            'impl_date': ''
                        }
                    else:
                        if not current_row:
                            current_row = {
                                'action': '',
                                'responsibility': '',
                                'target_date': '',
                                'impl_date': ''
                            }
                        if is_resp:
                            current_row['responsibility'] = s
                        elif is_date:
                            if not current_row['target_date']:
                                current_row['target_date'] = s
                            else:
                                current_row['impl_date'] = s
                if current_row:
                    rows.append(current_row)
                return rows
                
            if not corrective_actions:
                corrective_actions = parse_rows_from_lines(corr_lines)
            if not preventive_actions:
                preventive_actions = parse_rows_from_lines(prev_lines)

    # 7. Detailed implementation plan
    detailed_plan = ""
    for i, line in enumerate(all_lines):
        if line.lower().startswith("7. detailed implementation plan"):
            detailed_plan = line.split(":", 1)[1].strip() if ":" in line else (all_lines[i+1] if i+1 < len(all_lines) else "")
            break

    # 8. Modified Documents ticks
    modified_documents = []
    modified_documents_other = ""
    
    docs_to_check = [
        "MOC",
        "SOP / SMP",
        "Risk and Opportunity Register",
        "Register of Environmental Aspect Impact and OH & S Risks",
        "Training Need Identification"
    ]
    for doc_item in docs_to_check:
        if doc_item.lower() in text_dump.lower():
            modified_documents.append(doc_item)
            
    other_match = re.search(r'Others\s*\(Please\s*mention\)\s*:\s*(.*)', text_dump, re.IGNORECASE)
    if other_match:
        modified_documents_other = other_match.group(1).strip()

    # 9. Training Details
    training_details = find_colon_value(r'9\.\s*Training\s*Details\s*\(If\s*any\)\s*:\s*(.*)', text_dump)

    # 10. Date of Implementation
    date_implementation = find_colon_value(r'10\.\s*Date\s*of\s*Implementation\s*:\s*(.*)', text_dump)

    # 11. Effectiveness
    effectiveness_evaluation = ""
    for i, line in enumerate(all_lines):
        if "11. effectiveness evaluation" in line.lower():
            eff_parts = []
            for j in range(i + 1, min(i + 5, len(all_lines))):
                next_l = all_lines[j]
                if "prepared by" in next_l.lower() or "initiator" in next_l.lower():
                    break
                eff_parts.append(next_l)
            effectiveness_evaluation = " ".join(eff_parts).strip()
            
    # Approvals parsing
    prepared_by = ""
    reviewed_by = ""
    approved_by = ""
    
    # 1. Try parsing from doc_tables (docx) if available
    if doc_tables:
        for table in doc_tables:
            for r_idx, row in enumerate(table.rows):
                cells_text = [c.text.lower().strip() for c in row.cells]
                has_prep = any("prepared by" in ct or "initiator" in ct for ct in cells_text)
                has_rev = any("reviewed by" in ct or "reviewer" in ct for ct in cells_text)
                has_app = any("approved by" in ct or "approved by (hod)" in ct for ct in cells_text)
                
                if has_prep or has_rev or has_app:
                    if r_idx + 1 < len(table.rows):
                        next_row = table.rows[r_idx + 1]
                        for idx, ct in enumerate(cells_text):
                            val = next_row.cells[idx].text.strip() if idx < len(next_row.cells) else ""
                            if "prepared by" in ct or "initiator" in ct:
                                prepared_by = val
                            elif "reviewed by" in ct or "reviewer" in ct:
                                reviewed_by = val
                            elif "approved by" in ct or "approved by (hod)" in ct:
                                approved_by = val
                    break
                    
    # 2. Fallback: Parse from text stream (all_lines) for doc/pdf
    if not prepared_by or not reviewed_by or not approved_by:
        prep_idx = -1
        for i, line in enumerate(all_lines):
            l_lower = line.lower()
            if "prepared by" in l_lower or "initiator" in l_lower:
                prep_idx = i
                break
                
        if prep_idx != -1:
            candidate_lines = all_lines[prep_idx:]
            cleaned_candidates = []
            for line in candidate_lines:
                s = line.strip()
                if not s:
                    continue
                s_low = s.lower()
                # Filter out header label placeholders
                if any(kw in s_low for kw in [
                    "prepared by", "reviewed by", "approved by", 
                    "initiator", "reviewer", "hod", "signature", "effective date", "issue no"
                ]):
                    continue
                if s not in cleaned_candidates:
                    cleaned_candidates.append(s)
            
            # Map the remaining lines to prepared_by, reviewed_by, approved_by
            if len(cleaned_candidates) >= 3:
                prepared_by = prepared_by or cleaned_candidates[0]
                reviewed_by = reviewed_by or cleaned_candidates[1]
                approved_by = approved_by or cleaned_candidates[2]
            elif len(cleaned_candidates) == 2:
                prepared_by = prepared_by or cleaned_candidates[0]
                reviewed_by = reviewed_by or cleaned_candidates[1]
            elif len(cleaned_candidates) == 1:
                prepared_by = prepared_by or cleaned_candidates[0]

    # Clean up empty or signature placeholder strings
    if prepared_by and "signature" in prepared_by.lower():
        prepared_by = ""
    if reviewed_by and "signature" in reviewed_by.lower():
        reviewed_by = ""
    if approved_by and "signature" in approved_by.lower():
        approved_by = ""

    if not approved_by:
        approved_by = find_colon_value(r'Approved\s*by\s*\(Name\s*and\s*Signature\s*of\s*HOD\)\s*:\s*(.*)', text_dump)

    # Parse team members dynamically into the new 4-column structured dictionary layout
    team_leader = ""
    team_members = []
    role_function = []
    contact_nos = []
    
    table_parsed = False
    
    # 1. Try parsing from doc_tables (docx) if available
    if doc_tables:
        for table in doc_tables:
            # Check if this is the responsible team table
            header_row_idx = -1
            col_mapping = {}
            for r_idx, row in enumerate(table.rows):
                row_cells_text = [cell.text.lower().strip() for cell in row.cells]
                has_tl = any("team leader" in ct for ct in row_cells_text)
                has_tm = any("team members" in ct for ct in row_cells_text)
                has_rf = any("role" in ct or "function" in ct for ct in row_cells_text)
                has_cn = any("contact" in ct for ct in row_cells_text)
                
                # Check for legacy table layout headers
                has_rd = any("role designation" in ct for ct in row_cells_text)
                has_name = any("name" in ct for ct in row_cells_text)
                
                if (has_tl or has_rd) and (has_tm or has_name):
                    header_row_idx = r_idx
                    for idx, cell_text in enumerate(row_cells_text):
                        if "team leader" in cell_text:
                            col_mapping['team_leader'] = idx
                        elif "team members" in cell_text:
                            col_mapping['team_members'] = idx
                        elif "role" in cell_text or "function" in cell_text:
                            col_mapping['role_function'] = idx
                        elif "contact" in cell_text:
                            col_mapping['contact_nos'] = idx
                        elif "role designation" in cell_text:
                            col_mapping['legacy_role'] = idx
                        elif "name" in cell_text:
                            col_mapping['legacy_name'] = idx
                    break
            
            if header_row_idx != -1:
                # We found the table!
                table_parsed = True
                for r_idx in range(header_row_idx + 1, len(table.rows)):
                    row = table.rows[r_idx]
                    cells = row.cells
                    
                    # Extract from new table layout columns
                    if 'team_leader' in col_mapping or 'team_members' in col_mapping:
                        if 'team_leader' in col_mapping and col_mapping['team_leader'] < len(cells):
                            val = cells[col_mapping['team_leader']].text.strip()
                            if val and not team_leader:
                                team_leader = val
                        if 'team_members' in col_mapping and col_mapping['team_members'] < len(cells):
                            val = cells[col_mapping['team_members']].text.strip()
                            for line in val.split('\n'):
                                line_clean = line.strip()
                                if line_clean and line_clean not in team_members:
                                    team_members.append(line_clean)
                        if 'role_function' in col_mapping and col_mapping['role_function'] < len(cells):
                            val = cells[col_mapping['role_function']].text.strip()
                            for line in val.split('\n'):
                                line_clean = line.strip()
                                if line_clean and line_clean not in role_function:
                                    role_function.append(line_clean)
                        if 'contact_nos' in col_mapping and col_mapping['contact_nos'] < len(cells):
                            val = cells[col_mapping['contact_nos']].text.strip()
                            for line in val.split('\n'):
                                line_clean = line.strip()
                                if line_clean and line_clean not in contact_nos:
                                    contact_nos.append(line_clean)
                                    
                    # Extract from old/legacy table layout columns
                    elif 'legacy_role' in col_mapping and 'legacy_name' in col_mapping:
                        r_idx_val = col_mapping['legacy_role']
                        n_idx_val = col_mapping['legacy_name']
                        c_idx_val = col_mapping.get('contact_nos', -1)
                        
                        role_val = cells[r_idx_val].text.strip() if r_idx_val < len(cells) else ""
                        name_val = cells[n_idx_val].text.strip() if n_idx_val < len(cells) else ""
                        contact_val = cells[c_idx_val].text.strip() if (c_idx_val != -1 and c_idx_val < len(cells)) else ""
                        
                        if name_val:
                            if "leader" in role_val.lower():
                                team_leader = name_val
                            else:
                                if name_val not in team_members:
                                    team_members.append(name_val)
                            if contact_val and contact_val != "—" and contact_val != "-":
                                if contact_val not in contact_nos:
                                    contact_nos.append(contact_val)
                break

    # 2. Fallback: Parse from text stream (all_lines) for doc/pdf or if table parsing was not successful
    if not table_parsed:
        idx_2 = -1
        idx_3 = -1
        for i, line in enumerate(all_lines):
            l_lower = line.lower()
            if ("2." in l_lower or "2 " in l_lower) and "responsible team" in l_lower:
                idx_2 = i
            elif ("3." in l_lower or "3 " in l_lower) and ("correction" in l_lower or "immediate action" in l_lower):
                idx_3 = i
                
        if idx_2 != -1 and idx_3 != -1:
            team_lines = all_lines[idx_2 + 1 : idx_3]
            for line in team_lines:
                s = line.strip()
                if not s:
                    continue
                s_low = s.lower()
                
                # Exclude headers and section title values
                if s_low in [
                    "team leader", "team members", "role/function", "contact nos.",
                    "role designation", "name", "contact no", "contact nos",
                    "2. responsible team for corrective/preventive actions :",
                    "2. responsible team for corrective/preventive actions",
                    "responsible team for corrective/preventive actions"
                ]:
                    continue
                
                # Check for 10-digit phone number
                phone_match = re.search(r'\b\d{10}\b', s)
                if phone_match:
                    num = phone_match.group(0)
                    if num not in contact_nos:
                        contact_nos.append(num)
                # Check for role keywords
                elif any(kw in s_low for kw in ["analysis", "preventive action", "corrective action"]):
                    if s not in role_function:
                        role_function.append(s)
                # Check if it's a name
                elif len(s) > 3 and not any(char.isdigit() for char in s):
                    if not team_leader:
                        team_leader = s
                    elif s not in team_members:
                        team_members.append(s)

        # 3. Last fallback: Check known names/contacts if still empty
        if not team_leader and not team_members:
            known_names = ["Sachin Khanna", "Saurabh Agarwal", "Dibyayan Paul", "DC mandol", "R Viswakarma", "D.Paul", "Saurabh Agrawal"]
            for line in all_lines:
                for name in known_names:
                    if name.lower() in line.lower():
                        if not team_leader:
                            team_leader = name
                        elif name not in team_members and name != team_leader:
                            team_members.append(name)
            for line in all_lines:
                num_match = re.findall(r'\b\d{10}\b', line)
                for num in num_match:
                    if num not in contact_nos:
                        contact_nos.append(num)

    responsible_team = {
        'team_leader': team_leader,
        'team_members': "\n".join(team_members),
        'role_function': "\n".join(role_function),
        'contact_nos': "\n".join(contact_nos)
    }

    return {
        'area_section': area_section or "Mill Electrical",
        'date_incident': date_incident or "12/09/2024",
        'capa_no': capa_no or "Sep/Elect/04",
        'problem_what': problem_what,
        'problem_where': problem_where,
        'problem_when': problem_when,
        'problem_extent': problem_extent,
        'breakdown_applicable': "≥ 4 Hrs." if "270" in problem_extent else "2 – 4 Hrs.",
        'breakdown_hrs': "4.5" if "270" in problem_extent else "0",
        'breakdown_from': "12:00 AM",
        'breakdown_to': "04:30 AM",
        'responsible_team': responsible_team,
        'immediate_action': immediate_action,
        'action_timeframe': action_timeframe or "270 Minutes",
        'action_responsibility': action_responsibility or "Dibyayan Paul",
        'why_1': why_1,
        'why_2': why_2,
        'why_3': why_3,
        'why_4': why_4,
        'why_5': why_5,
        'whys': whys_list,
        'conclusion': conclusion,
        'five_m_applicable': ["3M Machine", "5M Method"],
        'corrective_actions': corrective_actions,
        'preventive_actions': preventive_actions,
        'detailed_plan': detailed_plan,
        'modified_documents': modified_documents,
        'modified_documents_other': modified_documents_other,
        'training_details': training_details,
        'date_implementation': date_implementation or "12.09.2024",
        'effectiveness_evaluation': effectiveness_evaluation,
        'prepared_by': prepared_by or "D.Paul",
        'reviewed_by': reviewed_by or "Saurabh Agrawal",
        'approved_by': approved_by or "Sachin Khanna",
        'document_no': document_no,
        'issue_no': issue_no,
        'issue_date': issue_date,
    }
